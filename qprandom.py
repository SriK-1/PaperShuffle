import os
import random
import ctypes
import re
import threading
import tkinter as tk
from tkinter import messagebox, filedialog
import customtkinter as ctk
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import openpyxl
import PyPDF2
from groq import Groq

# --- windows DPI fix for crisp UI on high-resolution displays ---
try:
    ctypes.windll.shcore.SetProcessDpiAwareness(1)
except Exception:
    pass

# --- apply dark mode theme globally ---
ctk.set_appearance_mode("Dark") 
ctk.set_default_color_theme("blue")

def get_set_name(index):
    """converts numbers to excel-style column names (0->A, 25->Z, 26->AA)"""
    result = ""
    while index >= 0:
        result = chr(65 + (index % 26)) + result
        index = (index // 26) - 1
    return result

def clean_question_number(q_text):
    """strips original question numbers (e.g., '1. ', 'Q1)') while preserving math decimals like '2.5'"""
    return re.sub(r'^\s*(?:[Qq]u?e?s?t?i?o?n?\s*)?\d+[\.\)\-](?:\s+|(?=[A-Za-z]))', '', q_text.strip())

def parse_questions_with_ai(raw_text, api_key):
    """sends text to groq api for semantic separation using llama-3"""
    try:
        client = Groq(api_key=api_key.strip())
        
        # to separate questions using *** delimiter
        prompt = """You are an expert exam parsing assistant. 
        I will give you raw text containing multiple exam questions.
        Your ONLY job is to separate the text into distinct, complete question blocks using exactly this delimiter: ***
        Keep all text within a question exactly as it is. Do NOT add any conversational text.
        Make sure sub-questions or instructions belong to their parent question block."""
        
        response = client.chat.completions.create(
            messages=[{"role": "system", "content": prompt}, {"role": "user", "content": raw_text}],
            model="llama3-70b-8192", 
            temperature=0.0  # deterministic output
        )
        
        ai_output = response.choices[0].message.content
        raw_questions = [q.strip() for q in ai_output.split('***') if q.strip()]
        
        # clean old question numbers from AI-separated text
        questions = []
        for q in raw_questions:
            cleaned_q = clean_question_number(q)
            if cleaned_q: questions.append(cleaned_q)
        return questions
    except Exception as e:
        raise Exception(f"AI Parsing Failed: {str(e)}")

def parse_questions_procedural(raw_text, mode):
    """local python regex parsing with three different strategies"""
    questions = []
    
    if mode == "Custom Separator (***)":
        # splits by user-defined ***
        raw_parts = raw_text.split('***')
        for q in raw_parts:
            cleaned_q = clean_question_number(q)
            if cleaned_q: questions.append(cleaned_q)
            
    elif mode == "Double Blank Lines":
        # splits by double newlines
        raw_parts = re.split(r'\n\s*\n', raw_text)
        for q in raw_parts:
            cleaned_q = clean_question_number(q)
            if cleaned_q: questions.append(cleaned_q)
            
    else:  # auto-detect numbers
        # regex finds lines starting with numbers like "1. " or "Q1)"
        pattern = re.compile(r'(?m)^(\s*(?:[Qq]u?e?s?t?i?o?n?\s*)?\d+[\.\)\-](?:\s+|(?=[A-Za-z])))')
        if pattern.search(raw_text):
            parts = pattern.split(raw_text)
            for i in range(1, len(parts), 2):
                combined = (parts[i] + parts[i+1]).strip()
                cleaned_q = clean_question_number(combined)
                if cleaned_q: questions.append(cleaned_q)
        else:
            # fallback to blank line separation if no numbers found
            questions = [q.strip() for q in re.split(r'\n\s*\n', raw_text) if q.strip()]
    return questions

class QuestionGeneratorApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # --- window configuration ---
        self.title("Pro Question Paper Generator")
        self.geometry("900x850")
        self.minsize(850, 700)  
        self.resizable(True, True)  

        # --- header ---
        self.title_label = ctk.CTkLabel(self, text="Question Paper Generator", font=ctk.CTkFont(size=26, weight="bold"))
        self.title_label.pack(pady=(15, 5))

        # --- tabs (upload files vs review text) ---
        self.tabview = ctk.CTkTabview(self)
        self.tabview.pack(padx=20, pady=5, fill="both", expand=True)
        self.tabview.add("Upload Files")
        self.tabview.add("Review & Edit Text")

        # --- tab 1: file upload interface ---
        self.filepaths = []  # stores selected file paths
        self.file_btn_frame = ctk.CTkFrame(self.tabview.tab("Upload Files"), fg_color="transparent")
        self.file_btn_frame.pack(fill="x", padx=10, pady=(10, 5))
        
        self.browse_btn = ctk.CTkButton(self.file_btn_frame, text="+ Add Files (Word, Excel, PDF)", command=self.browse_files)
        self.browse_btn.pack(side="left")
        
        self.clear_btn = ctk.CTkButton(self.file_btn_frame, text="Clear List", fg_color="#dc3545", hover_color="#c82333", command=self.clear_files)
        self.clear_btn.pack(side="right")
        
        self.file_listbox = ctk.CTkTextbox(self.tabview.tab("Upload Files"), state="disabled")
        self.file_listbox.pack(padx=10, pady=5, fill="both", expand=True)

        # --- tab 2: text preview/editing interface ---
        self.text_top_frame = ctk.CTkFrame(self.tabview.tab("Review & Edit Text"), fg_color="transparent")
        self.text_top_frame.pack(fill="x", padx=10, pady=(5, 0))
        
        self.text_instruction = ctk.CTkLabel(self.text_top_frame, text="Fix formatting mistakes here. Ensure distinct questions are separated by ***")
        self.text_instruction.pack(side="left")
        
        self.clear_text_btn = ctk.CTkButton(self.text_top_frame, text="Clear Text", width=80, height=24, fg_color="#dc3545", hover_color="#c82333", command=lambda: self.textbox.delete("1.0", tk.END))
        self.clear_text_btn.pack(side="right")

        self.textbox = ctk.CTkTextbox(self.tabview.tab("Review & Edit Text"))
        self.textbox.pack(padx=10, pady=5, fill="both", expand=True)

        # --- settings panel ---
        self.settings_frame = ctk.CTkScrollableFrame(self, height=250)
        self.settings_frame.pack(padx=20, pady=10, fill="x")

        # exam header input
        self.header_var = tk.StringVar(value="Midterm Examination")
        ctk.CTkLabel(self.settings_frame, text="Exam Header:", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, padx=10, pady=10, sticky="w")
        ctk.CTkEntry(self.settings_frame, textvariable=self.header_var, width=180).grid(row=0, column=1, padx=10, sticky="w")

        # output filename input
        self.filename_var = tk.StringVar(value="Question_Paper")
        ctk.CTkLabel(self.settings_frame, text="File Name:", font=ctk.CTkFont(weight="bold")).grid(row=0, column=2, padx=10, pady=10, sticky="w")
        ctk.CTkEntry(self.settings_frame, textvariable=self.filename_var, width=180).grid(row=0, column=3, padx=10, sticky="w")

        # number of sets to generate
        self.sets_var = tk.StringVar(value="3")
        ctk.CTkLabel(self.settings_frame, text="Number of Sets:", font=ctk.CTkFont(weight="bold")).grid(row=1, column=0, padx=10, pady=10, sticky="w")
        ctk.CTkEntry(self.settings_frame, textvariable=self.sets_var, width=60).grid(row=1, column=1, padx=10, sticky="w")

        # single file vs multiple files output
        self.output_format_var = tk.StringVar(value="Multiple Files")
        ctk.CTkLabel(self.settings_frame, text="Output Format:", font=ctk.CTkFont(weight="bold")).grid(row=1, column=2, padx=10, pady=10, sticky="w")
        ctk.CTkOptionMenu(self.settings_frame, variable=self.output_format_var, values=["Multiple Files", "One Single File"]).grid(row=1, column=3, padx=10, sticky="w")

        # parsing engine selection
        self.parsing_mode_var = tk.StringVar(value="Auto-Detect Numbers (1., Q1.)")
        ctk.CTkLabel(self.settings_frame, text="Parsing Engine:", font=ctk.CTkFont(weight="bold")).grid(row=2, column=0, padx=10, pady=10, sticky="w")
        self.parsing_dropdown = ctk.CTkOptionMenu(self.settings_frame, variable=self.parsing_mode_var, values=["Auto-Detect Numbers (1., Q1.)", "Double Blank Lines", "Custom Separator (***)", "AI Powered (Requires Groq API Key)"], width=280)
        self.parsing_dropdown.grid(row=2, column=1, columnspan=2, padx=10, sticky="w")

        # optional groq api key input
        self.api_key_var = tk.StringVar()
        ctk.CTkLabel(self.settings_frame, text="Groq API Key (Optional):", font=ctk.CTkFont(weight="bold"), text_color="gray").grid(row=3, column=0, padx=10, pady=10, sticky="w")
        self.api_entry = ctk.CTkEntry(self.settings_frame, textvariable=self.api_key_var, width=280, show="*")
        self.api_entry.grid(row=3, column=1, columnspan=2, padx=10, sticky="w")

        # toggle answer key inclusion
        self.answers_var = tk.BooleanVar(value=True)
        self.answer_checkbox = ctk.CTkCheckBox(self.settings_frame, text="Include Answer Key & Explanations in Output", variable=self.answers_var, font=ctk.CTkFont(weight="bold"))
        self.answer_checkbox.grid(row=4, column=0, columnspan=4, padx=10, pady=15, sticky="w")

        # --- action buttons (two-step workflow) ---
        self.btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.btn_frame.pack(pady=10)

        self.preview_btn = ctk.CTkButton(self.btn_frame, text="Step 1: Extract & Preview", font=ctk.CTkFont(weight="bold"), width=180, height=40, fg_color="#17a2b8", hover_color="#138496", command=self.start_preview_thread)
        self.preview_btn.pack(side="left", padx=10)

        self.generate_btn = ctk.CTkButton(self.btn_frame, text="Step 2: Generate Papers", font=ctk.CTkFont(weight="bold"), width=180, height=40, fg_color="#28a745", hover_color="#218838", command=self.start_generation_thread)
        self.generate_btn.pack(side="left", padx=10)

        # --- progress tracking ---
        self.progress_bar = ctk.CTkProgressBar(self)
        self.progress_bar.pack(padx=20, pady=(5, 5), fill="x")
        self.progress_bar.set(0)
        
        self.status_label = ctk.CTkLabel(self, text="Ready. Upload files and click Step 1.", text_color="gray")
        self.status_label.pack(pady=(0, 5))

    def browse_files(self):
        """opens file dialog and adds selected files to the list"""
        new_files = filedialog.askopenfilenames(filetypes=[("Supported Files", "*.txt *.docx *.xlsx *.pdf")])
        if new_files:
            self.filepaths.extend(new_files)
            self.filepaths = list(dict.fromkeys(self.filepaths))  # remove duplicates
            self.update_file_listbox()

    def clear_files(self):
        """removes all files from the upload list"""
        self.filepaths = []
        self.update_file_listbox()

    def update_file_listbox(self):
        """refreshes the file list display"""
        self.file_listbox.configure(state="normal")
        self.file_listbox.delete("1.0", tk.END)
        for i, path in enumerate(self.filepaths, 1):
            self.file_listbox.insert(tk.END, f"{i}. {os.path.basename(path)}\n")
        self.file_listbox.configure(state="disabled")

    def get_text_from_file(self, filepath):
        """extracts raw text from word, excel, pdf, or text files"""
        ext = os.path.splitext(filepath)[1].lower()
        try:
            if ext == '.txt':
                with open(filepath, 'r', encoding='utf-8') as f: 
                    return f.read()
            elif ext == '.docx':
                # extracts paragraph text from word doc
                return "\n".join([p.text for p in docx.Document(filepath).paragraphs])
            elif ext == '.xlsx':
                # reads all columns from excel and joins them
                lines = []
                for row in openpyxl.load_workbook(filepath, data_only=True).active.iter_rows(values_only=True):
                    r = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if r: lines.append("\n".join(r))
                return "\n\n".join(lines)
            elif ext == '.pdf':
                # extracts text from each pdf page
                text = ""
                for page in PyPDF2.PdfReader(filepath).pages:
                    extracted = page.extract_text()
                    if extracted: text += extracted + "\n\n"
                return text
            return ""
        except Exception as e:
            raise Exception(f"Error reading '{os.path.basename(filepath)}': {e}")

    # --- step 1: preview logic ---
    def start_preview_thread(self):
        """validates input and starts background parsing"""
        raw_text_input = self.textbox.get("1.0", tk.END).strip()
        if not self.filepaths and not raw_text_input:
            messagebox.showerror("Error", "Please upload files or paste text into the Review box first.")
            return
        
        # lock both buttons to prevent race conditions
        self.preview_btn.configure(state="disabled")
        self.generate_btn.configure(state="disabled")
        self.status_label.configure(text="Extracting and parsing...", text_color="yellow")
        threading.Thread(target=self.preview_logic, daemon=True).start()

    def preview_logic(self):
        """background worker that extracts and parses questions"""
        try:
            raw_text = ""
            
            # read from files or use pasted text
            if self.filepaths:
                for fp in self.filepaths:
                    raw_text += self.get_text_from_file(fp) + "\n\n"
            else:
                raw_text = self.textbox.get("1.0", tk.END).strip()
            
            # select parsing method    
            mode = self.parsing_mode_var.get()
            if "AI Powered" in mode:
                key = self.api_key_var.get()
                if not key: raise Exception("API Key required for AI Mode.")
                self.after(0, lambda: self.status_label.configure(text="AI is parsing questions... (Takes a few seconds)", text_color="yellow"))
                questions = parse_questions_with_ai(raw_text, key)
            else:
                questions = parse_questions_procedural(raw_text, mode)

            if len(questions) <= 1:
                raise Exception("Failed to separate multiple questions. Try a different Parsing Engine.")

            # inject *** separators between questions
            formatted_preview = "\n\n***\n\n".join(questions)
            
            # update UI
            self.after(0, lambda: self.textbox.delete("1.0", tk.END))
            self.after(0, lambda: self.textbox.insert(tk.END, formatted_preview))
            self.after(0, lambda: self.tabview.set("Review & Edit Text"))
            self.after(0, lambda: self.parsing_mode_var.set("Custom Separator (***)"))
            self.after(0, lambda: self.status_label.configure(text="Preview Ready! Please review the text and click Step 2.", text_color="#17a2b8"))
            
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", str(e)))
            self.after(0, lambda: self.status_label.configure(text="Error during Preview.", text_color="red"))
        finally:
            # unlock buttons
            self.after(0, lambda: self.preview_btn.configure(state="normal"))
            self.after(0, lambda: self.generate_btn.configure(state="normal"))

    # --- step 2: generation logic ---
    def start_generation_thread(self):
        """validates settings and starts background document generation"""
        raw_text = self.textbox.get("1.0", tk.END).strip()
        if not raw_text:
            messagebox.showerror("Error", "No text found! Please complete Step 1 or paste text into the Review tab.")
            return
        
        # validate number of sets    
        try:
            num_sets = int(self.sets_var.get().strip())
            if num_sets < 1: raise ValueError
        except:
            messagebox.showerror("Error", "Enter a valid positive number of sets.")
            return

        # ask user where to save
        save_dir = filedialog.askdirectory()
        if not save_dir: return

        self.generate_btn.configure(state="disabled")
        self.preview_btn.configure(state="disabled")
        self.progress_bar.set(0)
        self.status_label.configure(text="Generating papers...", text_color="yellow")
        
        mode = self.parsing_mode_var.get()
        out_format = self.output_format_var.get()
        header = self.header_var.get()
        base_name = self.filename_var.get()
        ans = self.answers_var.get()
        
        threading.Thread(target=self.generate_logic, args=(raw_text, num_sets, save_dir, mode, out_format, header, base_name, ans), daemon=True).start()

    def generate_logic(self, raw_text, num_sets, save_dir, mode, out_format, header_text, base_name, include_answers):
        """background worker that creates randomized word documents"""
        try:
            # auto-switch to custom separator if stars are detected
            if "***" in raw_text and mode != "Custom Separator (***)":
                mode = "Custom Separator (***)"
                self.after(0, lambda: self.parsing_mode_var.set("Custom Separator (***)"))

            # parse questions
            questions = parse_questions_procedural(raw_text, mode)
            if len(questions) <= 1: 
                raise Exception("Not enough questions found. Ensure they are separated properly.")
            
            safe_name = re.sub(r'[\\/*?:"<>|]', "", base_name).strip() or "Question_Paper"
            
            if len(os.path.join(save_dir, f"{safe_name}_Set_ZZZZZ.docx")) >= 250:
                raise Exception("The folder path is too long. Save closer to your C: drive.")

            # generate documents
            if out_format == "One Single File":
                doc = docx.Document()
                for i in range(num_sets):
                    self.after(0, lambda p=(i/num_sets): self.progress_bar.set(p))
                    random.shuffle(questions)  # randomize order
                    set_name = get_set_name(i)
                    self.add_custom_header(doc, set_name, header_text)
                    for idx, q in enumerate(questions, 1): 
                        self.write_question_to_doc(doc, idx, q, include_answers)
                    if i < num_sets - 1: 
                        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
                doc.save(os.path.join(save_dir, f"{safe_name}_All_Sets.docx"))
            else:
                # create separate files for each set
                for i in range(num_sets):
                    self.after(0, lambda p=(i/num_sets): self.progress_bar.set(p))
                    random.shuffle(questions)
                    set_name = get_set_name(i)
                    doc = docx.Document()
                    self.add_custom_header(doc, set_name, header_text)
                    for idx, q in enumerate(questions, 1): 
                        self.write_question_to_doc(doc, idx, q, include_answers)
                    doc.save(os.path.join(save_dir, f"{safe_name}_Set_{set_name}.docx"))

            # success notif
            self.after(0, lambda: self.progress_bar.set(1.0))
            self.after(0, lambda: self.status_label.configure(text="Finished Successfully", text_color="#28a745"))
            self.after(0, lambda: messagebox.showinfo("Success", f"Generated {num_sets} sets based on {len(questions)} questions in {save_dir}"))
            
        except PermissionError:
            self.after(0, lambda: messagebox.showerror("Error", "Permission Denied! Close the Word document if you have it open."))
            self.after(0, lambda: self.status_label.configure(text="Error: File locked.", text_color="red"))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", str(e)))
            self.after(0, lambda: self.status_label.configure(text="Error occurred.", text_color="red"))
        finally:
            self.after(0, lambda: self.generate_btn.configure(state="normal"))
            self.after(0, lambda: self.preview_btn.configure(state="normal"))

    def write_question_to_doc(self, doc, index, q_text, include_answers):
        """formats and writes a single question to the word document"""
        lines = q_text.split('\n')
        
        # write question number and first line
        doc.add_paragraph().add_run(f"Q{index}. {lines[0]}")
        
        # pattern to detect answer key lines
        ans_pattern = re.compile(r'^\s*(?:correct\s+)?ans(?:wer)?\s*(?:key)?\s*[:\-\=]', re.IGNORECASE)
        in_ans = False
        
        # write remaining lines with answer detection
        for line in lines[1:]:
            line_str = line.strip()
            if not line_str: continue
            
            if ans_pattern.match(line_str): 
                in_ans = True  # entered answer block
                
            if in_ans:
                if not include_answers: continue  # skip answer for student version
                # highlight answer in bold green
                run = doc.add_paragraph().add_run(line_str)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 150, 0)
            else:
                # normal text formatting
                doc.add_paragraph().add_run(line_str)
        
        # add spacing between questions
        doc.add_paragraph("")

    def add_custom_header(self, doc, set_name, header_text):
        """adds exam header and set identifier to the document"""
        heading = doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if header_text.strip() else WD_ALIGN_PARAGRAPH.RIGHT
        
        if header_text.strip():
            # add custom exam title
            run1 = heading.add_run(f"{header_text.strip()}\n")
            run1.font.size, run1.font.color.rgb = Pt(16), RGBColor(0, 0, 0)
        
        # add set identifier (A, B, C, etc.)
        run2 = heading.add_run(f"SET {set_name}")
        run2.font.size, run2.font.color.rgb = Pt(14), RGBColor(100, 100, 100)

if __name__ == "__main__":
    app = QuestionGeneratorApp()
    app.mainloop()
